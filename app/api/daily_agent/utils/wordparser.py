import os
from docx import Document

class WordParser:
    def __init__(self, chunk_size: int = 1500, chunk_overlap: int = 300):
        # Word는 텍스트 밀도가 PDF보다 낮을 수 있어 기본 청크 사이즈를 조금 더 크게 잡기도 합니다.
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def extract_text(self, file_path: str) -> list[str]:
        """Word 파일(.docx)에서 텍스트를 추출하고 청킹하여 반환"""
        if not os.path.exists(file_path):
            return []

        try:
            doc = Document(file_path)
            full_text = []

            # 1. 단락(Paragraph) 추출
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            # 2. 표(Table) 내부 텍스트 추출 (Word는 표 정보가 중요할 때가 많음)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))

            combined_text = "\n\n".join(full_text)

            if not combined_text.strip():
                return []

            # 3. 청킹 로직 (PDFParser와 동일한 방식 적용)
            return self._make_chunks(combined_text)

        except Exception as e:
            print(f"Word 파싱 에러: {e}")
            return []

    def _make_chunks(self, text: str) -> list[str]:
        """텍스트를 설정된 사이즈에 맞춰 조각냄"""
        chunks = []
        if len(text) <= self.chunk_size:
            return [text]

        start = 0
        while start < len(text):
            end = start + self.chunk_size
            chunks.append(text[start:end])
            start += (self.chunk_size - self.chunk_overlap)
            if end >= len(text):
                break
        return chunks